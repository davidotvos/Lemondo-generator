
# Import sys for getting the command line arguments
import sys
import os

# Import docx to work with .docx files.
# Must be installed: pip install python-docx
from docx import Document

from docx2pdf import convert
import datetime
import comtypes.client
import webbrowser



template_path = 'test_lemondo.docx'
save_folder = ''

# Lemondók listája
lemondoLi= []

testDict = {
    '_tervcim' : 'ABC',
    '_tipus' : 'Építési',
    '_iktatoszam' : 'Debrecen123',
    '_varos' : 'Debrecen',
    '_datum' : '2023.01.02',
    '_felelos' : 'Alföldi Imre',
    '_pozicio' : 'Debrecen Run Team Lead'
}

varos = 'Debrecen'
_felelos = 'Alföldi Imre'

honapok = {
    '01' : 'január',
    '02' : 'február',
    '03' : 'március',
    '04' : 'április',
    '05' : 'május',
    '06' : 'június',
    '07' : 'július',
    '08' : 'augusztus',
    '09' : 'szeptember',
    '10' : 'október',
    '11' : 'november',
    '12' : 'december'
}

class Lemondo:
    def __init__(self, tervcim, iktatoszam, tipus, datum):
        self.tervcim = tervcim
        self.iktatoszam = iktatoszam
        self.tipus = tipus
        self.dátum = datum
        self.felelos = _felelos

    def __str__(self):
        return f'tervcím: {self.tervcim}, iktatószám: {self.iktatoszam}, típus: {self.tipus}'

def create_pdfs():

    if len(lemondoLi) == 0:
        return

    doc = Document('test.docx')


    
    for l in lemondoLi:
        temp_dic = lemondo_to_dict(l)
        for key,value in temp_dic.items():
            docx_find_replace_text(doc, key, value)
        
        doc.save(f"{save_folder}/Lemondás - {temp_dic['_tervcim'].replace('/','-')}.docx")

    print(save_folder)
    webbrowser.open(save_folder)
    



"""
Ez a metódus a jelenlegi dátumot a "év. hónap. nap." formátumban adja vissza, 
ahol a hónapot a magyar név (a honapok szótárban megadva) jelzi. A datetime 
modult és egy szótárt (honapok) használja ehhez.
"""
def getdatum():
    d = str(datetime.date.today()).split('-')
    temp = honapok.get(d[1])
    d[1] = temp
    datum = f"{d[0]}. {d[1]} {d[2]}."
    return datum


"""
Ez a metódus egy adott állományt (input_path) PDF formátumba konvertál és elmenti 
egy megadott helyre (output_path). Az átalakítást a convert metódus végzi.
"""
def convert_to_pdf(input_path, output_path):
    convert(input_path, output_path)


"""
Ez a metódus a doc objektumban a search_text stringet a replace_text-re cseréli. 
Az átalakítás minden cellában és bekezdésben végbe megy, így biztosítva, hogy mindenhol 
megjelenő szöveg átalakításra kerüljön. Az átalakítás során a szöveg formázása is megmarad.
"""
def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text


def lemondo_to_dict(lemondo:Lemondo):
    
    d = {
        '_tervcim' : lemondo.tervcim,
        '_tipus' : lemondo.tipus,
        '_iktatoszam' : lemondo.iktatoszam,
        '_varos' : 'Debrecen',
        '_datum' : getdatum(),
        '_felelos' : 'Alföldi Imre',
        '_pozicio' : 'Debrecen Run Team Lead'
    }

    return d

